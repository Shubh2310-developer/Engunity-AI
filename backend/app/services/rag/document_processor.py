"""
Document Processor for RAG Pipeline
===================================

Comprehensive document processing service for ingesting, analyzing, and
preparing documents for the RAG system with BGE retrieval and Phi-2 generation.

Features:
- Multi-format document support (PDF, TXT, DOCX, MD, HTML)
- Intelligent text extraction and cleaning
- Document metadata extraction and enrichment
- Content chunking and preprocessing for embeddings
- Integration with BGE retriever for vectorization

Author: Engunity AI Team  
"""

import os
import json
import logging
import hashlib
import mimetypes
from typing import Dict, List, Optional, Any, Union, Tuple
from dataclasses import dataclass, asdict
from pathlib import Path
from datetime import datetime
import re

# Document processing libraries
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

# Local imports
from .bge_retriever import BGERetriever, DocumentChunk

logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Document metadata structure."""
    document_id: str
    filename: str
    file_path: str
    file_type: str
    file_size: int
    mime_type: str
    created_at: datetime
    processed_at: Optional[datetime] = None
    
    # Content metadata
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    
    # Processing metadata
    extraction_method: Optional[str] = None
    processing_status: str = "pending"
    error_message: Optional[str] = None
    
    # Content analysis
    language: str = "en"
    content_type: str = "document"  # document, code, manual, academic, etc.
    topics: List[str] = None
    
    def __post_init__(self):
        if self.topics is None:
            self.topics = []
    
    def to_dict(self) -> Dict[str, Any]:
        data = asdict(self)
        # Convert datetime objects to ISO strings
        for key, value in data.items():
            if isinstance(value, datetime):
                data[key] = value.isoformat()
        return data

@dataclass
class ProcessedDocument:
    """Processed document with extracted content and metadata."""
    metadata: DocumentMetadata
    content: str
    chunks: List[DocumentChunk]
    processing_stats: Dict[str, Any]

class DocumentProcessor:
    """Main document processing service."""
    
    def __init__(
        self,
        bge_retriever: Optional[BGERetriever] = None,
        storage_path: str = "./data/documents",
        chunk_size: int = 512,
        chunk_overlap: int = 50
    ):
        """
        Initialize document processor.
        
        Args:
            bge_retriever: BGE retriever instance for vectorization
            storage_path: Path to store processed documents
            chunk_size: Maximum tokens per chunk
            chunk_overlap: Overlap between chunks
        """
        self.bge_retriever = bge_retriever
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(parents=True, exist_ok=True)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Supported file types
        self.supported_types = {
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.html': self._process_html,
            '.htm': self._process_html,
            '.py': self._process_code,
            '.js': self._process_code,
            '.ts': self._process_code,
            '.java': self._process_code,
            '.cpp': self._process_code,
            '.c': self._process_code,
            '.json': self._process_json,
            '.xml': self._process_xml
        }
        
        logger.info("Document Processor initialized")
    
    def process_file(
        self, 
        file_path: Union[str, Path],
        document_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
        add_to_retriever: bool = True
    ) -> ProcessedDocument:
        """
        Process a single document file.
        
        Args:
            file_path: Path to document file
            document_id: Optional custom document ID
            metadata: Additional metadata
            add_to_retriever: Whether to add to BGE retriever
            
        Returns:
            Processed document
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Processing document: {file_path.name}")
        
        # Generate document ID if not provided
        if document_id is None:
            document_id = self._generate_document_id(file_path)
        
        # Create metadata
        doc_metadata = self._create_metadata(file_path, document_id, metadata)
        
        try:
            # Extract content based on file type
            content = self._extract_content(file_path, doc_metadata)
            
            # Process and clean content
            processed_content = self._clean_content(content, doc_metadata)
            
            # Analyze content
            self._analyze_content(processed_content, doc_metadata)
            
            # Create chunks
            chunks = self._create_chunks(processed_content, doc_metadata)
            
            # Update processing status
            doc_metadata.processed_at = datetime.now()
            doc_metadata.processing_status = "completed"
            doc_metadata.word_count = len(processed_content.split())
            doc_metadata.character_count = len(processed_content)
            
            # Processing stats
            processing_stats = {
                'chunks_created': len(chunks),
                'content_length': len(processed_content),
                'processing_time': (datetime.now() - doc_metadata.created_at).total_seconds(),
                'extraction_method': doc_metadata.extraction_method
            }
            
            # Create processed document
            processed_doc = ProcessedDocument(
                metadata=doc_metadata,
                content=processed_content,
                chunks=chunks,
                processing_stats=processing_stats
            )
            
            # Add to retriever if requested
            if add_to_retriever and self.bge_retriever:
                self.bge_retriever.add_document(
                    text=processed_content,
                    document_id=document_id,
                    metadata=doc_metadata.to_dict()
                )
            
            # Save processed document
            self._save_processed_document(processed_doc)
            
            logger.info(f"Successfully processed {file_path.name}: {len(chunks)} chunks created")
            return processed_doc
            
        except Exception as e:
            logger.error(f"Error processing {file_path.name}: {e}")
            doc_metadata.processing_status = "failed"
            doc_metadata.error_message = str(e)
            raise
    
    def process_directory(
        self,
        directory_path: Union[str, Path],
        recursive: bool = True,
        file_pattern: str = "*",
        add_to_retriever: bool = True
    ) -> List[ProcessedDocument]:
        """
        Process all supported documents in a directory.
        
        Args:
            directory_path: Path to directory
            recursive: Whether to process subdirectories
            file_pattern: File pattern to match
            add_to_retriever: Whether to add to BGE retriever
            
        Returns:
            List of processed documents
        """
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        logger.info(f"Processing directory: {directory_path}")
        
        # Find all files
        if recursive:
            files = list(directory_path.rglob(file_pattern))
        else:
            files = list(directory_path.glob(file_pattern))
        
        processed_docs = []
        
        for file_path in files:
            if file_path.suffix.lower() in self.supported_types:
                try:
                    processed_doc = self.process_file(
                        file_path=file_path,
                        add_to_retriever=add_to_retriever
                    )
                    processed_docs.append(processed_doc)
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    continue
        
        logger.info(f"Processed {len(processed_docs)} documents from {directory_path}")
        return processed_docs
    
    def _generate_document_id(self, file_path: Path) -> str:
        """Generate unique document ID from file path and content."""
        file_info = f"{file_path.name}_{file_path.stat().st_size}_{file_path.stat().st_mtime}"
        return hashlib.md5(file_info.encode()).hexdigest()
    
    def _create_metadata(
        self, 
        file_path: Path, 
        document_id: str,
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> DocumentMetadata:
        """Create document metadata."""
        stat_info = file_path.stat()
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        metadata = DocumentMetadata(
            document_id=document_id,
            filename=file_path.name,
            file_path=str(file_path),
            file_type=file_path.suffix.lower(),
            file_size=stat_info.st_size,
            mime_type=mime_type or "application/octet-stream",
            created_at=datetime.now()
        )
        
        # Add additional metadata if provided
        if additional_metadata:
            for key, value in additional_metadata.items():
                if hasattr(metadata, key):
                    setattr(metadata, key, value)
        
        return metadata
    
    def _extract_content(self, file_path: Path, metadata: DocumentMetadata) -> str:
        """Extract content from file based on type."""
        file_type = file_path.suffix.lower()
        
        if file_type in self.supported_types:
            extractor = self.supported_types[file_type]
            content = extractor(file_path, metadata)
            metadata.extraction_method = extractor.__name__
            return content
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _process_text(self, file_path: Path, metadata: DocumentMetadata) -> str:
        """Process plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'ascii']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"Could not decode text file: {file_path}")
    
    def _process_markdown(self, file_path: Path, metadata: DocumentMetadata) -> str:
        """Process Markdown files."""
        content = self._process_text(file_path, metadata)
        
        if MARKDOWN_AVAILABLE:
            # Convert to HTML then extract text
            html = markdown.markdown(content)
            if HTML_AVAILABLE:
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text()
        
        return content
    
    def _process_pdf(self, file_path: Path, metadata: DocumentMetadata) -> str:
        """Process PDF files."""
        if not PDF_AVAILABLE:
            raise RuntimeError("PDF processing libraries not available. Install PyPDF2 and pdfplumber.")
        
        content = ""
        
        # Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata.page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    metadata.page_count = len(pdf_reader.pages)
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n"
            except Exception as e2:
                logger.error(f"Both PDF extractors failed for {file_path}: {e2}")
                raise
        
        return content
    
    def _process_docx(self, file_path: Path, metadata: DocumentMetadata) -> str:
        """Process DOCX files."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("DOCX processing library not available. Install python-docx.")
        
        doc = DocxDocument(file_path)
        content = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))
        
        return "\n".join(content)
    
    def _process_html(self, file_path: Path, metadata: DocumentMetadata) -> str:
        """Process HTML files."""
        if not HTML_AVAILABLE:
            raise RuntimeError("HTML processing library not available. Install beautifulsoup4.")
        
        html_content = self._process_text(file_path, metadata)
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract title if available
        title_tag = soup.find('title')
        if title_tag:
            metadata.title = title_tag.get_text().strip()
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        return soup.get_text()
    
    def _process_code(self, file_path: Path, metadata: DocumentMetadata) -> str:
        """Process code files."""
        content = self._process_text(file_path, metadata)
        metadata.content_type = "code"
        
        # Add programming language info
        lang_map = {
            '.py': 'python',
            '.js': 'javascript', 
            '.ts': 'typescript',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c'
        }
        
        if file_path.suffix.lower() in lang_map:
            metadata.topics.append(lang_map[file_path.suffix.lower()])
        
        return content
    
    def _process_json(self, file_path: Path, metadata: DocumentMetadata) -> str:
        """Process JSON files."""
        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        # Convert JSON to readable text
        return json.dumps(json_data, indent=2)
    
    def _process_xml(self, file_path: Path, metadata: DocumentMetadata) -> str:
        """Process XML files.""" 
        if not HTML_AVAILABLE:
            return self._process_text(file_path, metadata)
        
        xml_content = self._process_text(file_path, metadata)
        soup = BeautifulSoup(xml_content, 'xml')
        return soup.get_text()
    
    def _clean_content(self, content: str, metadata: DocumentMetadata) -> str:
        """Clean and normalize extracted content."""
        # Remove excessive whitespace
        content = re.sub(r'\n\s*\n', '\n\n', content)
        content = re.sub(r' +', ' ', content)
        
        # Remove common PDF artifacts
        content = re.sub(r'\f', '\n', content)  # Form feed
        content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', content)  # Control chars
        
        # Normalize line endings
        content = content.replace('\r\n', '\n').replace('\r', '\n')
        
        return content.strip()
    
    def _analyze_content(self, content: str, metadata: DocumentMetadata):
        """Analyze content to extract additional metadata."""
        # Simple topic extraction based on keywords
        topic_keywords = {
            'programming': ['function', 'class', 'method', 'algorithm', 'code', 'programming'],
            'mathematics': ['equation', 'formula', 'theorem', 'proof', 'calculate'],
            'science': ['research', 'study', 'experiment', 'analysis', 'data'],
            'business': ['strategy', 'market', 'customer', 'revenue', 'business'],
            'technical': ['system', 'design', 'architecture', 'implementation', 'technical']
        }
        
        content_lower = content.lower()
        detected_topics = []
        
        for topic, keywords in topic_keywords.items():
            if any(keyword in content_lower for keyword in keywords):
                detected_topics.append(topic)
        
        metadata.topics.extend(detected_topics)
        
        # Try to extract title from content if not already set
        if not metadata.title:
            lines = content.split('\n')
            for line in lines[:10]:  # Check first 10 lines
                line = line.strip()
                if line and not line.startswith('#') and len(line) < 100:
                    metadata.title = line
                    break
    
    def _create_chunks(self, content: str, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Create document chunks for embedding."""
        if self.bge_retriever:
            # Use BGE retriever's chunking
            return self.bge_retriever.chunk_document(
                text=content,
                document_id=metadata.document_id,
                metadata=metadata.to_dict()
            )
        else:
            # Simple chunking fallback
            return self._simple_chunk(content, metadata)
    
    def _simple_chunk(self, content: str, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Simple text chunking when BGE retriever is not available."""
        chunks = []
        words = content.split()
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            if chunk_text.strip():
                chunk = DocumentChunk(
                    id=f"{metadata.document_id}_chunk_{len(chunks)}",
                    content=chunk_text,
                    document_id=metadata.document_id,
                    chunk_index=len(chunks),
                    metadata=metadata.to_dict()
                )
                chunks.append(chunk)
        
        return chunks
    
    def _save_processed_document(self, processed_doc: ProcessedDocument):
        """Save processed document to storage."""
        doc_dir = self.storage_path / processed_doc.metadata.document_id
        doc_dir.mkdir(exist_ok=True)
        
        # Save metadata
        with open(doc_dir / "metadata.json", 'w') as f:
            json.dump(processed_doc.metadata.to_dict(), f, indent=2)
        
        # Save content
        with open(doc_dir / "content.txt", 'w', encoding='utf-8') as f:
            f.write(processed_doc.content)
        
        # Save processing stats
        with open(doc_dir / "stats.json", 'w') as f:
            json.dump(processed_doc.processing_stats, f, indent=2)
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get overall processing statistics."""
        processed_docs = list(self.storage_path.glob("*/metadata.json"))
        
        stats = {
            'total_documents': len(processed_docs),
            'file_types': {},
            'content_types': {},
            'processing_status': {}
        }
        
        for metadata_path in processed_docs:
            try:
                with open(metadata_path, 'r') as f:
                    metadata = json.load(f)
                
                # Count file types
                file_type = metadata.get('file_type', 'unknown')
                stats['file_types'][file_type] = stats['file_types'].get(file_type, 0) + 1
                
                # Count content types
                content_type = metadata.get('content_type', 'document')
                stats['content_types'][content_type] = stats['content_types'].get(content_type, 0) + 1
                
                # Count processing status
                status = metadata.get('processing_status', 'unknown')
                stats['processing_status'][status] = stats['processing_status'].get(status, 0) + 1
                
            except Exception as e:
                logger.warning(f"Error reading metadata from {metadata_path}: {e}")
        
        return stats

# Factory function
def create_document_processor(bge_retriever: Optional[BGERetriever] = None, **kwargs) -> DocumentProcessor:
    """Create document processor with default configuration."""
    return DocumentProcessor(bge_retriever=bge_retriever, **kwargs)

# Export main classes
__all__ = [
    "DocumentProcessor",
    "DocumentMetadata",
    "ProcessedDocument",
    "create_document_processor"
]